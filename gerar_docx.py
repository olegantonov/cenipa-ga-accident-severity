"""
Converte manuscrito.md em um .docx com formatacao de submissao (Elsevier-like):
Times New Roman 12, corpo justificado em espaco duplo, numeracao de linhas
continua (para revisao), titulos numerados, tabelas Word nativas, figuras
embutidas com legendas, referencias com recuo deslocado, e paginacao.

Uso: .venv/bin/python gerar_docx.py
Saida: manuscrito.docx
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).parent
MD = BASE / "manuscrito.md"
OUT = BASE / "manuscrito.docx"
FIGDIR = BASE / "figuras"
FONT = "Times New Roman"
SENT = "\x00"  # sentinela para asterisco literal

FIG_CAPTIONS = {
    "fig1_letalidade_faixa.png": "Figure 1. Crude proportion of fatal accidents by pilot-in-command total flight-hour band (error bars: 95% confidence intervals).",
    "fig2_prob_ajustada.png": "Figure 2. Adjusted predicted probability of a fatal accident versus PIC total flight hours (natural cubic spline), by type of operation.",
    "fig3_forest.png": "Figure 3. Adjusted odds ratios (95% confidence intervals) for a fatal accident from the multivariable logistic model.",
    "fig4_simpson.png": "Figure 4. Median PIC total flight hours by type of operation and outcome, illustrating confounding (Simpson's paradox).",
    "fig5_tipo_vs_total.png": "Figure 5. Adjusted predicted probability of a fatal accident versus in-type versus total flight hours.",
}


def set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)


def add_line_numbers(section):
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "continuous")
    ln.set(qn("w:distance"), "360")
    section._sectPr.append(ln)


def add_page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    a = OxmlElement("w:fldChar"); a.set(qn("w:fldCharType"), "begin")
    b = OxmlElement("w:instrText"); b.set(qn("xml:space"), "preserve"); b.text = "PAGE"
    c = OxmlElement("w:fldChar"); c.set(qn("w:fldCharType"), "end")
    run._r.append(a); run._r.append(b); run._r.append(c)
    run.font.name = FONT; run.font.size = Pt(10)


def add_runs(paragraph, text, base_size=12):
    text = text.replace("\\" + "*", SENT)
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        bold = ital = False
        if part.startswith("**") and part.endswith("**"):
            part = part[2:-2]; bold = True
        elif part.startswith("*") and part.endswith("*"):
            part = part[1:-1]; ital = True
        run = paragraph.add_run(part.replace(SENT, "*"))
        run.font.name = FONT; run.font.size = Pt(base_size)
        run.bold = bold; run.italic = ital


def style_heading(p, before=12, after=6):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.keep_with_next = True


def add_table(doc, rows):
    header = [c.strip() for c in rows[0].strip("|").split("|")]
    body = [[c.strip() for c in r.strip("|").split("|")] for r in rows[2:]]
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]; cell.paragraphs[0].text = ""
        add_runs(cell.paragraphs[0], "**" + h + "**", base_size=10)
    for r in body:
        cells = t.add_row().cells
        for j, val in enumerate(r):
            if j < len(cells):
                cells[j].paragraphs[0].text = ""
                add_runs(cells[j].paragraphs[0], val, base_size=10)
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def main():
    lines = MD.read_text(encoding="utf-8").split("\n")
    doc = Document()
    set_base_style(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.5)
        s.left_margin = s.right_margin = Cm(2.5)
        add_line_numbers(s)
        add_page_number_footer(s)

    i = 0
    in_refs = False
    seen_h2 = False
    while i < len(lines):
        s = lines[i].strip()
        if not s or s == "---":
            i += 1; continue

        if s.startswith("# ") and not s.startswith("## "):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(s[2:]); r.bold = True; r.font.name = FONT; r.font.size = Pt(15)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            i += 1; continue

        if s.startswith("### "):
            p = doc.add_paragraph(); style_heading(p)
            r = p.add_run(s[4:]); r.bold = True; r.italic = True; r.font.name = FONT; r.font.size = Pt(12)
            i += 1; continue

        if s.startswith("## "):
            seen_h2 = True
            title = s[3:]
            in_refs = title.lower().startswith("references")
            p = doc.add_paragraph(); style_heading(p)
            r = p.add_run(title); r.bold = True; r.font.name = FONT; r.font.size = Pt(13)
            i += 1; continue

        if s.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip()); i += 1
            if len(block) >= 2:
                add_table(doc, block)
            continue

        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, s[2:])
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(2)
            i += 1; continue

        if s.startswith("*Figures:") or s.startswith("*Figure files"):
            i += 1; continue

        if in_refs and re.match(r"^\[\d+\]", s):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.75)
            p.paragraph_format.space_after = Pt(4)
            add_runs(p, s, base_size=11)
            i += 1; continue

        # parágrafo comum; bloco de autoria (antes do 1º '## ') centralizado
        p = doc.add_paragraph()
        add_runs(p, s)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if not seen_h2 else WD_ALIGN_PARAGRAPH.JUSTIFY
        i += 1

    # Figuras
    doc.add_page_break()
    h = doc.add_paragraph(); style_heading(h)
    rr = h.add_run("Figures"); rr.bold = True; rr.font.name = FONT; rr.font.size = Pt(13)
    nf = 0
    for fname, cap in FIG_CAPTIONS.items():
        fp = FIGDIR / fname
        if not fp.exists():
            continue
        nf += 1
        pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pic.add_run().add_picture(str(fp), width=Cm(15))
        cp = doc.add_paragraph()
        cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cp.paragraph_format.space_after = Pt(10)
        add_runs(cp, cap, base_size=10)

    doc.save(OUT)
    print(f"OK -> {OUT.name} | tabelas: {len(doc.tables)} | figuras: {nf} | parágrafos: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
